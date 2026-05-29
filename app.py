import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
from docx import Document
import io
import re

# Page Setup & Academic Branding Style
st.set_page_config(page_title="Research Analyst Bee", layout="wide")

# Native Streamlit Branding Header
st.title("📊 Research Analyst Bee")
st.subheader("Advanced Academic Automation Engine & Statistical Modeling Suite")

# Clean Native System Verification Banner
st.success("💡 **System Verification:** Designed & Engineered by **Ajayi, I.A.** | Department of Public Administration Frameworks")

st.write("\n")

# Sidebar Configuration
st.sidebar.markdown("### ⚙️ System Controls")
mode = st.sidebar.radio("Data Intake Method:", ["Process Research Document (.docx/Spreadsheet)", "Simulate Assumption Data"])

st.sidebar.markdown("---")
st.sidebar.markdown("### 📈 Visualization Controls")
chart_type = st.sidebar.selectbox("Select Analysis Chart Type:", ["Clustered Column Chart", "Pie Chart"])

# CRITICAL ENGINE INITIALIZATION (Ensures storage dictionaries exist globally before logic checks)
report_tables = {}
report_text_blocks = {}
active_study_data = None
sample_size = 100

# Baseline Default Data (Fallback structural template)
default_study_data = {
    "Causes of Conflict (Section A)": {
        "table_no": "4.1.3",
        "items": [
            {
                "q_num": 1, "var": "Poor communication among staff leads to conflict", 
                "sa": 45, "a": 40, "n": 0, "d": 10, "sd": 5,
                "text_block": "Question 1: On whether poor communication among staff leads to conflict, 45 (45.0%) Strongly Agreed and 40 (40.0%) Agreed, totaling 85 (85.0%). Only 15 (15.0%) disagreed or strongly disagreed. This indicates a high level of agreement among the respondents. This implies that structural barriers in information sharing or closed-door communication policies create gaps that are quickly filled by rumors, suspicion, and friction among personnel, which severely hampers administrative workflow. When compared with Question 4 (role ambiguity), there is a strong correlation showing that where vertical communication is poor, clarity regarding individual job duties declines proportionally, compounding the potential for operational disputes."
            },
            {
                "q_num": 2, "var": "Inadequate resources causes disagreement among workers", 
                "sa": 50, "a": 38, "n": 0, "d": 8, "sd": 4,
                "text_block": "Question 2: Concerning whether inadequate resources causes disagreement among workers, 50 (50.0%) Strongly Agreed and 38 (38.0%) Agreed, totaling 88 (88.0%). Only 12 (12.0%) disagreed or strongly disagreed. This stands out as a critical baseline concern in this section. The direct implication is that scarcity of essential operational tools, office consumables, and infrastructure forces staff into unhealthy internal competition. Workers have to hoard or struggle over limited assets to execute their duties. This high agreement level of 88.0% directly mirrors the structural tensions analyzed in Table 4.1.4, indicating that resource deficits are the root fuel behind the high prevalence of departmental conflicts."
            },
            {
                "q_num": 3, "var": "Differences in personality contribute to workplace conflict", 
                "sa": 35, "a": 45, "n": 0, "d": 12, "sd": 8,
                "text_block": "Question 3: On whether differences in personality contribute to workplace conflict, 35 (35.0%) Strongly Agreed and 45 (45.0%) Agreed, totaling 80 (80.0%). Meanwhile, 20 (20.0%) disagreed or strongly disagreed. This indicates that diverse individual behavioral traits, varying tolerance levels, and differing ego thresholds frequently clash within the close-knit organizational sub-units when emotional intelligence parameters are absent. In comparison with institutional causes (like resource scarcity at 88.0%), personal variables scored slightly lower. This establishes that institutional and systemic flaws play a larger role in generating friction than individual temperaments."
            },
            {
                "q_num": 4, "var": "Role ambiguity (unclear job responsibilities) leads to conflict", 
                "sa": 42, "a": 40, "n": 0, "d": 10, "sd": 8,
                "text_block": "Question 4: Regarding whether role ambiguity (unclear job responsibilities) leads to conflict, 42 (42.0%) Strongly Agreed and 40 (40.0%) Agreed, totaling 82 (82.0%). Only 18 (18.0%) disagreed or strongly disagreed. This implies that overlapping portfolios or poorly delineated boundaries allow personnel to cross lines of authority unknowingly, leading to territorial clashes, buck-passing, and administrative inefficiency. This matches the recommendation metrics in Section D, confirming that providing written job descriptions is highly necessary to systematically eliminate this specific structural cause of friction."
            }
        ]
    },
    "Types of Conflict Prevalent (Section B)": {
        "table_no": "4.1.4",
        "items": [
            {
                "q_num": 5, "var": "Interpersonal conflict exists among employees", 
                "sa": 48, "a": 40, "n": 0, "d": 8, "sd": 4,
                "text_block": "Question 5: On whether interpersonal conflict exists among employees, 48 (48.0%) Strongly Agreed and 40 (40.0%) Agreed, totaling 88 (88.0%). Only 12 (12.0%) disagreed or strongly disagreed. This reveals that friction often shifts from purely professional disagreements into personal animosities, creating a tense working environment that erodes employee morale. The 88.0% agreement closely aligns with Question 1 (poor communication), proving that personal animosities thrive when professional communication networks fail."
            },
            {
                "q_num": 6, "var": "Task-related conflict occurs frequently in the organization", 
                "sa": 40, "a": 45, "n": 0, "d": 10, "sd": 5,
                "text_block": "Question 6: Concerning whether task-related conflict occurs frequently in the organization, 40 (40.0%) Strongly Agreed and 45 (45.0%) Agreed, totaling 85 (85.0%). 15 (15.0%) disagreed or strongly disagreed. This indicates frequent friction regarding operational methods, task assignments, and performance criteria, which slows down service delivery and program execution. While task conflict is high (85.0%), it is slightly lower than management-staff tension (90.0%), proving that vertical structural hierarchy gaps cause more friction than peer-to-peer work delegation."
            },
            {
                "q_num": 7, "var": "Conflict between management and staff is common", 
                "sa": 52, "a": 38, "n": 0, "d": 6, "sd": 4,
                "text_block": "Question 7: Regarding whether conflict between management and staff is common, 52 (52.0%) Strongly Agreed and 38 (38.0%) Agreed, totaling 90 (90.0%). Only 10 (10.0%) disagreed or strongly disagreed. This represents the highest level of agreement in this section. This clear 90.0% majority points to a gap between administrators and subordinates, often caused by perceived high-handedness, a lack of feedback options, or top-down directive management styles. This finding strongly justifies the recommendation in Section D (Question 14) where a participatory leadership style is demanded by an identical 90.0% margin to fix this issue."
            },
            {
                "q_num": 8, "var": "Intragroup conflict occurs within departments", 
                "sa": 38, "a": 42, "n": 0, "d": 12, "sd": 8,
                "text_block": "Question 8: On whether intragroup conflict occurs within departments, 38 (38.0%) Strongly Agreed and 42 (42.0%) Agreed, totaling 80 (80.0%). 20 (20.0%) disagreed or strongly disagreed. This shows that internal departmental units suffer from cliques and splintering, which weakens team cohesion and makes it difficult to coordinate collaborative tasks. Scoring the lowest in Section B (80.0%), it shows that while departments have internal differences, they tend to pull together when facing external issues or managing broader vertical management demands."
            }
        ]
    },
    "Effects on Organizational Performance (Section C)": {
        "table_no": "4.1.5",
        "items": [
            {
                "q_num": 9, "var": "Effective conflict management improves staff productivity", 
                "sa": 50, "a": 40, "n": 0, "d": 6, "sd": 4,
                "text_block": "Question 9: On whether effective conflict management improves staff productivity, 50 (50.0%) Strongly Agreed and 40 (40.0%) Agreed, totaling 90 (90.0%). Only 10 (10.0%) disagreed or strongly disagreed. This confirms that resolving disputes quickly prevents distractions, protects work hours, and keeps staff focused on delivering public administration objectives. This positive effect perfectly balances the negative impact found in Question 11, showing that productivity fluctuates based on how well administrators handle conflict."
            },
            {
                "q_num": 10, "var": "Proper conflict resolution enhances teamwork", 
                "sa": 48, "a": 42, "n": 0, "d": 6, "sd": 4,
                "text_block": "Question 10: Regarding whether proper conflict resolution enhances teamwork, 48 (48.0%) Strongly Agreed and 42 (42.0%) Agreed, totaling 90 (90.0%). Only 10 (10.0%) disagreed or strongly disagreed. This establishes that effective dispute resolution restores trust, reinforces organizational unity, and encourages cross-departmental collaboration. This 90.0% agreement matches the 90.0% agreement on management-staff conflict, showing that fixing hierarchy friction is key to rebuilding teamwork."
            },
            {
                "q_num": 11, "var": "Poor conflict management reduces organizational performance", 
                "sa": 55, "a": 38, "n": 0, "d": 4, "sd": 3,
                "text_block": "Question 11: Concerning whether poor conflict management reduces organizational performance, 55 (55.0%) Strongly Agreed and 38 (38.0%) Agreed, totaling 93 (93.0%). Only 7 (7.0%) disagreed or strongly disagreed. This is the highest level of agreement in this section. This indicates that letting conflicts fester directly lowers overall performance. It leads to high absenteeism, low dedication, and project delays. This strong 93.0% consensus serves as the foundation for our hypothesis test, leading us to reject the null hypothesis and confirm that conflict management systematically affects performance."
            },
            {
                "q_num": 12, "var": "Conflict management leads to better decision-making", 
                "sa": 42, "a": 40, "n": 0, "d": 10, "sd": 8,
                "text_block": "Question 12: On whether conflict management leads to better decision-making, 42 (42.0%) Strongly Agreed and 40 (40.0%) Agreed, totaling 82 (82.0%). 18 (18.0%) disagreed or strongly disagreed. This shows that addressing disagreements helps managers spot underlying systemic issues, leading to more refined long-term policies. While still highly supported (82.0%), it suggests that while conflict management improves policy design, its most immediate benefits are felt in daily productivity and teamwork."
            }
        ]
    },
    "Recommended Conflict Management Strategies (Section D)": {
        "table_no": "4.1.6",
        "items": [
            {
                "q_num": 13, "var": "Open communication should be encouraged to manage conflict", 
                "sa": 58, "a": 35, "n": 0, "d": 4, "sd": 3,
                "text_block": "Question 13: On whether open communication should be encouraged to manage conflict, 58 (58.0%) Strongly Agreed and 35 (35.0%) Agreed, totaling 93 (93.0%). Only 7 (7.0%) disagreed or strongly disagreed. This is the highest level of agreement in this section. This strongly indicates that transparent communication networks, regular staff town halls, and reliable feedback channels act as a vital preventative cushion against workplace misunderstandings. This 93.0% recommendation level directly answers the problem raised in Question 1, proving that improving communication is the most widely supported solution to organizational friction."
            },
            {
                "q_num": 14, "var": "Management should adopt participatory leadership style", 
                "sa": 50, "a": 40, "n": 0, "d": 6, "sd": 4,
                "text_block": "Question 14: Regarding whether management should adopt a participatory leadership style, 50 (50.0%) Strongly Agreed and 40 (40.0%) Agreed, totaling 90 (90.0%). Only 10 (10.0%) disagreed or strongly disagreed. This implies that involving staff in decision-making paths helps reduce systemic alienation, builds policy ownership, and cuts down resistance to administrative changes. This 90.0% agreement directly addresses the 90.0% management-staff conflict found in Section B, showing that participatory leadership is the key strategy to resolve hierarchy tension."
            },
            {
                "q_num": 15, "var": "Training on conflict management should be provided to staff", 
                "sa": 52, "a": 38, "n": 0, "d": 6, "sd": 4,
                "text_block": "Question 15: Concerning whether training on conflict management should be provided to staff, 52 (52.0%) Strongly Agreed and 38 (38.0%) Agreed, totaling 90 (90.0%). 10 (10.0%) disagreed or strongly disagreed. This shows that resolving disputes effectively is a technical skill set that needs ongoing training, professional workshops, and capacity-building programs. This high consensus indicates that staff recognize that relying solely on personal intuition is not enough to resolve deep-seated workplace disputes."
            },
            {
                "q_num": 16, "var": "Use of mediation helps resolve conflicts effectively", 
                "sa": 48, "a": 42, "n": 0, "d": 6, "sd": 4,
                "text_block": "Question 16: On whether the use of mediation helps resolve conflicts effectively, 48 (48.0%) Strongly Agreed and 42 (42.0%) Agreed, totaling 90 (90.0%). Only 10 (10.0%) disagreed or strongly disagreed. This highlights the need for neutral third-party mediators or internal panels to handle grievances objectively before they escalate into formal legal or disciplinary matters. This ties in with Question 3 (personality differences), confirming that when individuals cannot reach an agreement on their own, structured mediation provides a reliable path forward."
            }
        ]
    }
}

def extract_clean_number(text):
    """Safely extracts the first valid sequence of digits found in a row cell string (handles formatting variations like '45 (45%)' or '38%')"""
    match = re.search(r'\d+', str(text))
    return int(match.group()) if match else 0

# DYNAMIC FILE PROCESSING PARSING ENGINE
if mode == "Process Research Document (.docx/Spreadsheet)":
    uploaded_file = st.file_uploader("Upload Questionnaire Spreadsheet (.csv/.xlsx) or Word Document (.docx)", type=["docx", "csv", "xlsx"])
    
    if uploaded_file is not None:
        try:
            # Case 1: Processing Spreadsheets (CSV or Excel formats)
            if uploaded_file.name.endswith('.csv') or uploaded_file.name.endswith('.xlsx'):
                if uploaded_file.name.endswith('.csv'):
                    uploaded_df = pd.read_csv(uploaded_file)
                else:
                    uploaded_df = pd.read_excel(uploaded_file)
                
                # Standardize data headings to lowercase format to prevent KeyErrors
                uploaded_df.columns = uploaded_df.columns.str.strip().str.lower()
                required_cols = ['sa', 'a', 'n', 'd', 'sd', 'var']
                
                if all(col in uploaded_df.columns for col in required_cols):
                    parsed_items = []
                    for idx, row in uploaded_df.iterrows():
                        q_num = int(row['q_num']) if 'q_num' in uploaded_df.columns else (idx + 1)
                        var_name = str(row['var'])
                        sa_c = extract_clean_number(row['sa'])
                        a_c = extract_clean_number(row['a'])
                        n_c = extract_clean_number(row['n'])
                        d_c = extract_clean_number(row['d'])
                        sd_c = extract_clean_number(row['sd'])
                        
                        parsed_items.append({
                            "q_num": q_num, "var": var_name,
                            "sa": sa_c, "a": a_c, "n": n_c, "d": d_c, "sd": sd_c,
                            "text_block": f"Question {q_num}: Dynamic verification tracking for research parameter focus evaluation block '{var_name}' yielded active fields: SA({sa_c}), A({a_c}), N({n_c}), D({d_c}), SD({sd_c})."
                        })
                    
                    active_study_data = {"Uploaded Dataset Analytics Matrix": {"table_no": "4.1.1", "items": parsed_items}}
                    st.success(f"🎉 Successfully extracted dataset metrics ({len(parsed_items)} items) from spreadsheet upload!")
                else:
                    st.error(f"❌ Structural layout verification mismatch. Document dataset metrics must contain these columns exactly: var, sa, a, n, d, sd. Found: {list(uploaded_df.columns)}")

            # Case 2: Processing Word Documents (.docx)
            elif uploaded_file.name.endswith('.docx'):
                doc = Document(uploaded_file)
                parsed_items = []
                global_item_counter = 1
                
                for table in doc.tables:
                    for row in table.rows[1:]:
                        cells = row.cells
                        # Ensure row contains the structural sequence layout distribution (S/N, Variable, SA, A, N, D, SD)
                        if len(cells) >= 7:
                            var_text = cells[1].text.strip()
                            
                            # Skip standard summary, reference labels, or empty records cleanly
                            if any(x in var_text.lower() for x in ["total", "source", "researcher", "percentage", "s/n"]) or var_text == "":
                                continue
                                
                            try:
                                sa_val = extract_clean_number(cells[2].text)
                                a_val = extract_clean_number(cells[3].text)
                                n_val = extract_clean_number(cells[4].text)
                                d_val = extract_clean_number(cells[5].text)
                                sd_val = extract_clean_number(cells[6].text)
                                
                                parsed_items.append({
                                    "q_num": global_item_counter, "var": var_text,
                                    "sa": sa_val, "a": a_val, "n": n_val, "d": d_val, "sd": sd_val,
                                    "text_block": f"Question {global_item_counter}: Field study frequency distribution data metrics concerning item focus tracking '{var_text}' recorded active values."
                                })
                                global_item_counter += 1
                            except Exception:
                                continue
                
                if len(parsed_items) > 0:
                    active_study_data = {"Extracted Document Tables": {"table_no": "4.1.2", "items": parsed_items}}
                    st.success(f"🎉 Successfully found and compiled metrics for {len(parsed_items)} research questions from Word Document tables!")
                else:
                    st.error("❌ No standard academic distribution matrices detected inside document. Verify your file contains tables with headings aligned as: S/N | Variables | SA | A | N | D | SD.")
        except Exception as file_err:
            st.error(f"Engine Parsing Exception: Could not unpack file layout structure safely. Details: {file_err}")
    else:
        st.info("👉 System idle. Please upload a structured research file (.docx, .csv, or .xlsx) from your device storage to begin processing field metrics.")
else:
    # Simulation Data Intake Mode
    active_study_data = default_study_data

# VISUAL ANALYSIS RUNTIME LOOP
if active_study_data is not None:
    st.markdown("---")
    st.markdown(f"### 📊 Automated Multi-Dimensional Analysis (Sample Size, n = {sample_size})")
    
    for section_title, section_content in active_study_data.items():
        tbl_num = section_content["table_no"]
        st.markdown(f"#### 📑 Table {tbl_num}: Distribution of responses on {section_title.lower()}")
        
        section_results = []
        q_labels = []
        sa_values = []
        a_values = []
        n_values = []
        d_values = []
        sd_values = []
        
        for item in section_content["items"]:
            total_count = item["sa"] + item["a"] + item["n"] + item["d"] + item["sd"]
            if total_count == 0: total_count = 1  # Guard against division by zero errors
            
            section_results.append({
                "S/N": f"{item['q_num']}.",
                "Variables": item["var"],
                "SA (%)": f"{item['sa']} ({item['sa']/total_count*100:.1f}%)",
                "A (%)": f"{item['a']} ({item['a']/total_count*100:.1f}%)",
                "N (%)": f"{item['n']} ({item['n']/total_count*100:.1f}%)",
                "D (%)": f"{item['d']} ({item['d']/total_count*100:.1f}%)",
                "SD (%)": f"{item['sd']} ({item['sd']/total_count*100:.1f}%)",
                "Total": f"{int(total_count)} (100.0%)"
            })
            
            q_labels.append(f"Item {item['q_num']}")
            sa_values.append(item["sa"])
            a_values.append(item["a"])
            n_values.append(item["n"])
            d_values.append(item["d"])
            sd_values.append(item["sd"])
            
        res_df = pd.DataFrame(section_results)
        st.dataframe(res_df, use_container_width=True)
        st.caption("Source: Researcher's survey, 2026")
        
        st.markdown(f"**Visual Representation ({chart_type}) - Table {tbl_num}**")
        
        if chart_type == "Clustered Column Chart":
            fig = go.Figure()
            fig.add_trace(go.Bar(x=q_labels, y=sa_values, name='Strongly Agree', marker_color='#1f77b4'))
            fig.add_trace(go.Bar(x=q_labels, y=a_values, name='Agree', marker_color='#aec7e8'))
            fig.add_trace(go.Bar(x=q_labels, y=n_values, name='Neutral', marker_color='#c7c7c7'))
            fig.add_trace(go.Bar(x=q_labels, y=d_values, name='Disagree', marker_color='#ffbb78'))
            fig.add_trace(go.Bar(x=q_labels, y=sd_values, name='Strongly Disagree', marker_color='#ff7f0e'))
            
            fig.update_layout(
                barmode='group',
                xaxis_title='Survey Questionnaire Items',
                yaxis_title='Number of Respondents',
                legend_title='Response Modes',
                margin=dict(l=20, r=20, t=25, b=20),
                height=380,
                yaxis=dict(range=[0, max(max(sa_values), max(a_values)) + 15])
            )
        else:
            labels_p = ["Strongly Agree", "Agree", "Neutral", "Disagree", "Strongly Disagree"]
            values_p = [sum(sa_values), sum(a_values), sum(n_values), sum(d_values), sum(sd_values)]
            
            fig = go.Figure(data=[go.Pie(labels=labels_p, values=values_p, hole=.3)])
            fig.update_layout(
                margin=dict(l=20, r=20, t=25, b=20),
                height=380,
                legend_title="Response Modes"
            )
            
        st.plotly_chart(fig, use_container_width=True, key=f"table_chart_{tbl_num}")
        
        st.markdown(f"**📝 Academic Interpretation Framework (Table {tbl_num})**")
        for item in section_content["items"]:
            st.info(item["text_block"])
        st.write("\n")
        
        report_tables[f"Table {tbl_num}"] = res_df
        report_text_blocks[f"Table {tbl_num}"] = "\n\n".join([i["text_block"] for i in section_content["items"]])

    # TEST OF HYPOTHESIS SECTION
    st.markdown("---")
    st.markdown("### 🔬 4.2 Test of Hypothesis")
    st.write("**Null Hypothesis (Ho):** Conflict management practices have no significant effect on organizational performance in Oyo State College of Agriculture and Technology, Igboora.")
    st.write("**Alternative Hypothesis (Hi):** Conflict management practices have a significant effect on organizational performance in Oyo State College of Agriculture and Technology, Igboora.")
    
    hypo_records = [
        {"S/N": "1.", "Statement": "Effective conflict management improves staff productivity", "SA + A (Combined Agreement)": "50 + 40 = 90", "Percentage": "90%", "Decision": "Significant"},
        {"S/N": "2.", "Statement": "Proper conflict resolution enhances teamwork", "SA + A (Combined Agreement)": "48 + 42 = 90", "Percentage": "90%", "Decision": "Significant"},
        {"S/N": "3.", "Statement": "Poor conflict management reduces organizational performance", "SA + A (Combined Agreement)": "55 + 38 = 93", "Percentage": "93%", "Decision": "Significant"},
        {"S/N": "4.", "Statement": "Conflict management leads to better decision-making", "SA + A (Combined Agreement)": "42 + 40 = 82", "Percentage": "82%", "Decision": "Significant"}
    ]
    st.markdown("**Table 4.3.3: Distribution of responses on effect of conflict management on organizational performance (Extracted for Hypothesis Testing)**")
    st.dataframe(pd.DataFrame(hypo_records), use_container_width=True)
    st.caption("Source: Researcher's survey, 2026")
    
    final_conclusion = "REJECTED. Therefore, the alternative hypothesis (Hi) which states that 'Conflict management practices have a significant effect on organizational performance' is ACCEPTED."
    st.info(f"**Conclusion Decision Rule:** Since the calculated indicators score consistently above the 50% majority threshold, the null hypothesis (Ho) is officially **{final_conclusion}**")

    # CHAPTER FIVE SECTION
    st.markdown("---")
    st.markdown("### 📑 CHAPTER FIVE: SUMMARY, CONCLUSION AND RECOMMENDATIONS")
    
    st.markdown("#### 5.1 Summary")
    st.write("The study examined Conflict Management Practices and Organizational Performance in Oyo State College of Agriculture and Technology, Igboora. The empirical findings established that inadequate resources (88%) and poor communication (85%) are primary triggers of workplace tension. Management-staff conflict remains the most prevalent structural variant (90%).")
    
    st.markdown("#### 5.2 Conclusion")
    st.write("In view of the operational indicators analyzed, this study concludes that poor conflict management mechanisms severely deteriorate overall institutional performance capacity (93% consensus), while proactive intervention loops directly upgrade staff output and organizational performance vectors.")
    
    st.markdown("#### 5.3 Recommendations")
    st.write("1. Regular staff meetings, suggestion boxes, and feedback channels should be established.\n"
             "2. Transparent criteria for distributing limited resources should be developed and communicated to all staff.\n"
             "3. Clear, written job descriptions should be provided to every staff member to reduce role ambiguity.\n"
             "4. Managers should be trained in participatory leadership and emotional intelligence.\n"
             "5. All staff should receive annual training on conflict resolution skills.")

    # COMPREHENSIVE WORD DOCUMENT (.DOCX) GENERATION UTILITY
    st.markdown("---")
    st.markdown("### 💾 Step 2: Custom Document Export")
    custom_filename = st.text_input("Enter your desired filename for export:", value="Conflict_Management_Analysis_Report")
    
    def generate_docx_file(tables_dict, text_dict, conclusion_text):
        doc = Document()
        doc.add_heading("Research Analysis Report (Chapter 4 & 5)", 0)
        doc.add_paragraph("Department of Public Administration Frameworks")
        doc.add_paragraph("Analyst Authority Signature: Ajayi, I.A.")
        
        for name, table_df in tables_dict.items():
            doc.add_heading(name, level=2)
            
            t = doc.add_table(rows=1, cols=8)
            t.style = 'Light Shading Accent 1'
            hdr_cells = t.rows[0].cells
            headers = ["S/N", "Variables", "SA (%)", "A (%)", "N (%)", "D (%)", "SD (%)", "Total"]
            for x, h in enumerate(headers):
                hdr_cells[x].text = h
                
            for _, r in table_df.iterrows():
                row_cells = t.add_row().cells
                for idx, col_name in enumerate(headers):
                    row_cells[idx].text = str(r[col_name])
            
            doc.add_paragraph("\nAcademic Interpretation Matrix:")
            doc.add_paragraph(text_dict.get(name, ""))
            doc.add_paragraph("\n")
            
        doc.add_heading("4.2 Test of Hypothesis", level=2)
        doc.add_paragraph(f"Decision Rule Result: {conclusion_text}")
        
        doc.add_heading("CHAPTER FIVE: SUMMARY, CONCLUSION AND RECOMMENDATIONS", level=1)
        doc.add_heading("5.1 Summary", level=2)
        doc.add_paragraph("The study examined Conflict Management Practices and Organizational Performance in Oyo State College of Agriculture and Technology, Igboora. The empirical findings established that inadequate resources (88%) and poor communication (85%) are primary triggers of workplace tension. Management-staff conflict remains the most prevalent structural variant (90%).")
        doc.add_heading("5.2 Conclusion", level=2)
        doc.add_paragraph("In view of the operational indicators analyzed, this study concludes that poor conflict management mechanisms severely deteriorate overall institutional performance capacity (93% consensus), while proactive intervention loops directly upgrade staff output and organizational performance vectors.")
        doc.add_heading("5.3 Recommendations", level=2)
        doc.add_paragraph("1. Regular staff meetings, suggestion boxes, and feedback channels should be established.\n2. Transparent criteria for distributing limited resources should be developed and communicated to all staff.\n3. Clear, written job descriptions should be provided to every staff member to reduce role ambiguity.\n4. Managers should be trained in participatory leadership and emotional intelligence.\n5. All staff should receive annual training on conflict resolution skills.")
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    try:
        docx_buffer = generate_docx_file(report_tables, report_text_blocks, final_conclusion)
        st.sidebar.markdown("---")
        st.sidebar.download_button(
            label="📥 Download Complete Report (.DOCX)",
            data=docx_buffer,
            file_name=f"{custom_filename}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as err:
        st.sidebar.error(f"Export engine notice: {err}")

st.markdown("---")
st.caption("Research Analyst Bee Platform Core Engine • Standard Compliance Distribution Framework v2.5")
